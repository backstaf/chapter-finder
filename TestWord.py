import os
import glob
import docx
import math
import string
import numpy as np

""" Letar reda på kapitel och paragrafer ur ett word-dokument från söktermer med hjälp av tf-idf. 
    Förbättringsförslag:
    * Spara ner vocabulary och tf-idf till fil för att slippa räkna om dem varje gång, och implementera en trigger för när de ska räknas om.
    * Ta med även tabeller och inte bara paragrafer
    * Prioritera avsnitt som har den exakta sökfrasen på lämpligt sätt
"""

def is_heading2(paragraph):
    """ Checks whether the paragraph contains a Heading
    """
    if paragraph.style.name.startswith('Heading 1') or paragraph.style.name.startswith('Heading 2'):
    # if paragraph.style.name.startswith('Heading'):
        return True
    else:
        return False

def iterate_document_sections(document):
    """Generate a sequence of paragraphs for each headed section in document.

    Each generated sequence has a heading paragraph in its first position, 
    followed by one or more body paragraphs.
    """
    paragraphs = [document.paragraphs[0]]
    for paragraph in document.paragraphs[1:]:
        if is_heading2(paragraph):
             yield paragraphs
             paragraphs = [paragraph]
             continue
        paragraphs.append(paragraph)
    yield paragraphs


def create_vocabulary_from_doc(document):
    """ Generate and return a vocabulary of all the words in the document paragraphs. 
    """
    all_paras = document.paragraphs
    parasplit = []
    for para in all_paras:
        para = para.text.replace("-", " ").replace("/", " ").translate(str.maketrans('', '', string.punctuation))
        parasp = para.lower().split()
        if (len(parasp)>0):
            parasplit.extend(parasp)

    vocab = list(set(parasplit))
    return vocab

doc = docx.Document("Manual.docx")
# print(doc.tables) doc.tables contains all the tables in the doc. Improvement could be to loop through and put these in the vocabulary and chapters
vocabulary = create_vocabulary_from_doc(doc)

paras = []

for paragraphs in iterate_document_sections(doc):
   paras.append(paragraphs)
""" Paras är en 2D array, där första dimensionen är delkapitel och andra paragrafer. Objekten är paragraph-objekt
    som behöver kallas med .text för att få ut strängen
"""

df = {}
tf = {}
searchdict = {}


# Nedan gör om paras till en lista av strängar som var och en räknas som ett dokument
docs = []
para = ""

for par in range(len(paras)): # par är kapitelnummer
    for pa in range(len(paras[par])): # pa är paragrafnummer
        para += " " + paras[par][pa].text.replace("-", " ").replace("/", " ").translate(str.maketrans('', '', string.punctuation)).lower()
    if len(para.split()) > 0:
        docs.append(para.split())
        para = ""

N = len(docs)

# Nedan är vad sökorden är som ska matchas mot dokumenten
searchterm = """lib.conf""" 
searchtermlist = searchterm.replace("-", " ").replace("/", " ").translate(str.maketrans('-', ' ', string.punctuation)).lower().split()

# Nedan tar fram tf och df för alla ord i vocabulary och för varje doc. Även en tf för searchdict
for word in vocabulary:
    tf[word] = [doc.count(word)/len(doc) for doc in docs] # tf är en 2D dict, där nyckeln är vocab-ord och värden arrayer av kapitel. Objekten är float som beskriver hur stor del av orden i aktuellt doc som är aktuellt ord.
    df[word] = sum([word in doc for doc in docs])/N # df är en 1D dict, som per vocab-ord beskriver i hur stor del av docsen ordet finns.
    searchdict[word] = searchtermlist.count(word)/len(searchtermlist) # searchworddict är en 1D dict, där nyckel är vocab-ord och värdet är hur stor del av orden i aktuellt doc som är aktuellt ord

# Nedan kombinerar tf och df till tf-idf, en för dokumenten och en för sökorden
tottfidf = [] # tottfidf är en 2D array där första dimensionen är vocab-ord och andra är docs. Innehåller tfidf-värden för dessa. 
tfidfsearch = [] # tfidfsearch är en 1D array som innehåller tfidf-värden per vocab-ord. 
tfidf = []

for doc_index in range(len(docs)):   
    tfidf.clear()
    for word in vocabulary:
        tfidf.append(tf[word][doc_index]*math.log(1/df[word],10)) 
        if doc_index == 0:
            tfidfsearch.append(searchdict[word]*math.log(1/df[word],10))              
    tottfidf.append(tfidf.copy())

def find_nearest_chapter(data, searchterms):
    """ Find and print the best tf-idf matches in the data based on the search terms.
    """

    bestresults = 3
    M = len(data)
    dist = np.empty(M, dtype=float)

    for rowcomp in range(M):
        dist[rowcomp] = 0
        for rowpos in range(len(data[rowcomp])):
            if searchterms[rowpos] > 0: # Only check the words that are in searchterms
                dist[rowcomp] += abs(data[rowcomp][rowpos]-searchterms[rowpos])
                # print(data[rowcomp][rowpos], rowcomp, searchterms[rowpos], dist[rowcomp])
        if dist[rowcomp] == 0 : dist[rowcomp] = np.inf
#    print(dist)

    chapters = np.argsort(dist)[:bestresults] # Find the chapters with the lowest total distance from the search terms

    # print (chapters)
    for index in chapters:
        print("\n\nCHAPTER INDEX: ", index, "MATCH: %.2f" % ((1/(1+dist[index]))*100), "%")
        for paragraphs in range(len(paras[index])):
            print (paras[index][paragraphs].text)

find_nearest_chapter(tottfidf, tfidfsearch)


